"""
DOCX Builder — Assembles extracted content into a formatted Word document.
"""

import io
import logging
import os

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image as PILImage

from .models import PageContent, TextResult, TextBlock, DocumentContent

logger = logging.getLogger(__name__)

FONT_SIZE_MAP = {
    "heading1": Pt(24), "heading2": Pt(18), "heading3": Pt(14),
    "body": Pt(11), "small": Pt(9),
}


def _map_font_size(estimated_size: float) -> Pt:
    if estimated_size >= 22:
        return FONT_SIZE_MAP["heading1"]
    elif estimated_size >= 16:
        return FONT_SIZE_MAP["heading2"]
    elif estimated_size >= 13:
        return FONT_SIZE_MAP["heading3"]
    elif estimated_size >= 9:
        return FONT_SIZE_MAP["body"]
    else:
        return FONT_SIZE_MAP["small"]


def _add_header_footer(doc: Document, header_text: TextResult | None, footer_text: TextResult | None):
    section = doc.sections[0]
    if header_text and not header_text.is_empty:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.text = header_text.raw_text.replace("\n\n", " | ")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0] if p.runs else p.add_run()
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)
        logger.debug(f"  Added header: {header_text.raw_text[:50]}...")
    if footer_text and not footer_text.is_empty:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = footer_text.raw_text.replace("\n\n", " | ")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0] if p.runs else p.add_run()
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)
        logger.debug(f"  Added footer: {footer_text.raw_text[:50]}...")


def _add_text_blocks(doc: Document, text_result: TextResult):
    if not text_result or text_result.is_empty:
        return
    for block in text_result.blocks:
        if not block.text.strip():
            continue
        if block.is_heading:
            if block.font_size_estimate >= 22:
                doc.add_heading(block.text, level=1)
            elif block.font_size_estimate >= 16:
                doc.add_heading(block.text, level=2)
            else:
                doc.add_heading(block.text, level=3)
        else:
            para = doc.add_paragraph()
            run = para.add_run(block.text)
            run.font.size = _map_font_size(block.font_size_estimate)
            if block.is_bold:
                run.bold = True


def _add_image(doc: Document, pil_image: PILImage.Image, max_width_inches: float = 6.0):
    try:
        img_buffer = io.BytesIO()
        pil_image.save(img_buffer, format="PNG")
        img_buffer.seek(0)
        img_width, img_height = pil_image.size
        width = Inches(min(max_width_inches, img_width / 96.0))
        doc.add_picture(img_buffer, width=width)
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logger.debug(f"  Added image ({img_width}x{img_height})")
    except Exception as e:
        logger.warning(f"  Failed to add image: {e}")


def build_document(document_content: DocumentContent, output_path: str):
    doc = Document()
    logger.info(f"Building document with {document_content.page_count} pages")

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    header_text = document_content.common_header
    footer_text = document_content.common_footer
    if not header_text and document_content.pages:
        header_text = document_content.pages[0].header_text
    if not footer_text and document_content.pages:
        footer_text = document_content.pages[0].footer_text
    _add_header_footer(doc, header_text, footer_text)

    for i, page in enumerate(document_content.pages):
        logger.info(f"  Processing page {page.page_number} ({i+1}/{document_content.page_count})")
        if page.body_text and not page.body_text.is_empty:
            _add_text_blocks(doc, page.body_text)
        for img in page.images:
            _add_image(doc, img)
        if i < document_content.page_count - 1:
            doc.add_page_break()

    out_dir = os.path.dirname(output_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)
    doc.save(output_path)
    logger.info(f"Document saved to: {output_path}")
